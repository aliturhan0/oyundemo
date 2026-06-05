from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    doc.add_paragraph()

# ===== TITLE =====
title = doc.add_heading('Proje İlerleme Raporu', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Low Poly Shooter – Zombi Hayatta Kalma (Survival)')
r.bold = True
r.font.size = Pt(14)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Motor: Unity 2022.3.62f3 (LTS) | Render Pipeline: URP 14.0.12\n').font.size = Pt(10)
info.add_run('Tarih: 11 Mayıs 2026').font.size = Pt(10)

doc.add_page_break()

# ===== 1. PROJE OZETI =====
add_heading('1. Proje Özeti', 1)
doc.add_paragraph(
    'Bu proje, Unity 2022.3 LTS üzerinde geliştirilen bir birinci şahıs (FPS) zombi hayatta kalma oyunudur. '
    'Oyuncu, low-poly stilinde tasarlanmış bir çevrede zombilere karşı savaşır. Oyunun temel döngüsü '
    '"Horde Mode" (Dalga Modu) üzerine kurulmuş olup, oyuncunun belirli sayıda zombiyi öldürerek bölümü geçmesi gerekmektedir.'
)
doc.add_paragraph(
    'Proje, Infima Games – Low Poly Shooter Pack (Free Sample) asset paketi üzerine inşa edilmiş ve üzerine '
    'özel zombi AI, spawner, oyuncu sağlık sistemi, ölüm/kazanma mekanikleri gibi sistemler eklenmiştir.'
)

# ===== 2. ASSET PAKETLERI =====
add_heading('2. Kullanılan Asset Paketleri', 1)
add_table(
    ['#', 'Asset Paketi', 'Açıklama', 'Kaynak'],
    [
        ['1', 'Infima Games – Low Poly Shooter Pack (Free Sample)', 'FPS karakter, silah sistemi, UI, animasyonlar ve ses efektleri', 'Unity Asset Store'],
        ['2', 'StarterAssets – First Person Controller', 'Birinci şahıs karakter kontrolcüsü (hareket + kamera)', 'Unity Asset Store'],
        ['3', 'NatureStarterKit2', 'Doğa ortamı (ağaç, çimen, yer kaplamaları, terrain materyalleri)', 'Unity Asset Store'],
        ['4', 'Real Stars Skybox', 'Gerçekçi yıldızlı gece gökyüzü skybox\'ı', 'Unity Asset Store'],
        ['5', 'Zombi Modeli (Mixamo – Ch10_nonPBR)', 'Zombi 3D karakteri ve animasyonları', 'Mixamo'],
    ]
)

# ===== 3. OZEL SCRIPTLER =====
add_heading('3. Eklenen / Yazılan Script\'ler (Özel Kodlar)', 1)
doc.add_paragraph(
    'Aşağıdaki script\'ler sıfırdan yazılmış veya projeye özel olarak eklenmiştir. '
    'Bunlar paketin içinde gelen kodlardan ayrıdır ve projenin özgün geliştirme kısmını oluşturur.'
)

add_heading('3.1 EnemyAI.cs — Zombi Yapay Zekası', 2)
doc.add_paragraph('Konum: Assets/Infima Games/Low Poly Shooter Pack - Free Sample/Code/')
doc.add_paragraph('Satır Sayısı: 216 satır')
bullet_items = [
    'NavMeshAgent ile otomatik yol bulma (pathfinding)',
    'Oyuncuya doğru hareket etme',
    'Saldırı menzili kontrolü ve hasar verme',
    'Vurulma tepkisi (hit stun) ve ölüm sistemi',
    'Ortam sesleri (hırlama), saldırı sesleri, vurulma sesleri',
    'Animator ile yürüme, saldırı, vurulma, ölüm animasyonları',
    'GameManager ve ZombieSpawner entegrasyonu',
    'Ölünce collider kapatma ve 5 saniye sonra destroy',
]
for item in bullet_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading('3.2 PlayerHealth.cs — Oyuncu Sağlık Sistemi', 2)
doc.add_paragraph('Konum: Assets/Infima Games/Low Poly Shooter Pack - Free Sample/Code/')
doc.add_paragraph('Satır Sayısı: 202 satır')
for item in [
    'Can sistemi (100 HP, UI Slider ile gösterim)',
    'Hasar alma ve kısa süreli dokunulmazlık (invincibility)',
    'Kanlı ekran efekti (blood screen overlay, fade in/out)',
    'Kalp atışı sesi (can %30\'un altına düşünce)',
    'Ölüm sekansı: Tüm kontroller devre dışı, kamera yere düşer, kan ekranı kızarır, ölüm menüsü açılır',
    'Time.timeScale ile oyun duraklatma',
    'Mouse kilit yönetimi (lock/unlock)',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading('3.3 GameManager.cs — Oyun Yöneticisi', 2)
doc.add_paragraph('Konum: Assets/Infima Games/Low Poly Shooter Pack - Free Sample/Code/')
doc.add_paragraph('Satır Sayısı: 105 satır')
for item in [
    'Singleton pattern',
    'Hedef: 15 zombi öldürerek bölüm geçme',
    'UI\'da canlı sayaç: "Öldürülen Zombi: X / 15"',
    'Kazanma durumunda: mesaj, ses ve kazanma menüsü',
    'Sonraki bölüme geçiş desteği',
    'TextMeshPro ile UI metin yönetimi',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading('3.4 ZombieSpawner.cs — Zombi Doğuş Sistemi', 2)
doc.add_paragraph('Konum: Assets/Infima Games/Low Poly Shooter Pack - Free Sample/Code/')
doc.add_paragraph('Satır Sayısı: 58 satır')
for item in [
    'Birden fazla spawn noktası desteği',
    'Ayarlanabilir spawn aralığı (varsayılan: 3 saniye)',
    'Eşzamanlı zombi limiti (varsayılan: max 5)',
    'Coroutine tabanlı sürekli spawn döngüsü',
    'EnemyAI ile entegrasyon (zombi ölünce slot açılır)',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading('3.5 Flashlight.cs — El Feneri Sistemi', 2)
doc.add_paragraph('Konum: Assets/Infima Games/Low Poly Shooter Pack - Free Sample/Code/')
doc.add_paragraph('Satır Sayısı: 39 satır')
for item in ['F tuşu ile açma/kapama (toggle)', 'Opsiyonel tıklama ses efekti', 'Light component kontrolü']:
    doc.add_paragraph(item, style='List Bullet')

add_heading('3.6 Collectible.cs — Eşya Toplama (Devre Dışı)', 2)
doc.add_paragraph('Satır Sayısı: 8 satır — Horde Mode\'a geçildiği için devre dışı bırakılmış.')

add_heading('3.7 SilahSistemi.cs — Silah Sistemi (Ek)', 2)
doc.add_paragraph('Konum: Assets/ (root)')
doc.add_paragraph('Satır Sayısı: ~80 satır')
for item in [
    'Menzil, mermi sayısı (30 mermi)',
    'Sol tık ile ateş etme',
    'R tuşu ile şarjör değiştirme (Coroutine + animasyon)',
    'Raycast tabanlı vuruş sistemi',
    'Muzzle flash (namlu ateşi) particle efekti',
    'Silah sesi ve animasyon entegrasyonu',
]:
    doc.add_paragraph(item, style='List Bullet')

# ===== 4. PAKET SCRIPTLERI =====
add_heading('4. Mevcut Asset Pack Script\'leri', 1)

add_heading('Karakter Sistemi (Code/Character/)', 2)
add_table(['Script', 'Açıklama'], [
    ['Character.cs', 'Ana karakter davranış sınıfı (20.8 KB)'],
    ['CharacterBehaviour.cs', 'Karakter abstract base class'],
    ['CharacterKinematics.cs', 'Karakter kinematik hesaplamaları'],
    ['Movement.cs', 'Hareket sistemi'],
    ['MovementBehaviour.cs', 'Hareket abstract class'],
    ['Inventory.cs', 'Silah envanteri yönetimi'],
    ['InventoryBehaviour.cs', 'Envanter abstract class'],
])

add_heading('Silah Sistemi (Code/Weapons/)', 2)
add_table(['Script', 'Açıklama'], [
    ['Weapon.cs', 'Ana silah sınıfı'],
    ['WeaponBehaviour.cs', 'Silah abstract class'],
    ['WeaponAttachmentManager.cs', 'Eklenti yönetimi'],
    ['Magazine.cs / MagazineBehaviour.cs', 'Şarjör sistemi'],
    ['Muzzle.cs / MuzzleBehaviour.cs', 'Namlu ateşi ve mermi çıkışı'],
    ['Scope.cs / ScopeBehaviour.cs', 'Nişangah/dürbün sistemi'],
    ['UtilitiesWeapons.cs', 'Silah yardımcı fonksiyonları'],
])

add_heading('Arayüz - UI (Code/Interface/)', 2)
add_table(['Script', 'Açıklama'], [
    ['CanvasSpawner.cs', 'Canvas oluşturucu'],
    ['Crosshair.cs', 'Nişangah (crosshair)'],
    ['ImageWeapon.cs', 'Silah ikonu gösterimi'],
    ['TextAmmunitionCurrent.cs', 'Mevcut mermi sayısı'],
    ['TextAmmunitionTotal.cs', 'Toplam mermi sayısı'],
    ['Element.cs / ElementText.cs', 'UI element yardımcıları'],
])

add_heading('Servis Sistemi (Code/Services/)', 2)
add_table(['Script', 'Açıklama'], [
    ['ServiceLocator.cs', 'Servis bulucu (dependency injection)'],
    ['GameModeService.cs / IGameModeService.cs', 'Oyun modu servisi'],
    ['AudioManagerService.cs / IAudioManagerService.cs', 'Ses yönetim servisi'],
    ['Bootstraper.cs', 'Başlatıcı'],
])

add_heading('Animasyon, Kamera ve Diğer', 2)
add_table(['Script', 'Açıklama'], [
    ['CharacterAnimationEventHandler.cs', 'Karakter animasyon olayları'],
    ['WeaponAnimationEventHandler.cs', 'Silah animasyon olayları'],
    ['PlaySoundBehaviour.cs', 'Ses çalma davranışı'],
    ['CameraLook.cs', 'Kamera bakış sistemi'],
    ['TimeHandler.cs', 'Zaman yönetimi'],
    ['FirstPersonController.cs', 'Birinci şahıs hareket kontrolcüsü (StarterAssets)'],
])

# ===== 5. SES DOSYALARI =====
add_heading('5. Ses Dosyaları (Audio Assets)', 1)
add_heading('Özel Eklenen Sesler (Assets/zombi/ses/)', 2)
add_table(['Dosya', 'Açıklama', 'Boyut'], [
    ['hirlama.m4a', 'Zombi hırlama (ortam sesi)', '254 KB'],
    ['zombihirilti.wav', 'Zombi hırıltısı', '1.35 MB'],
    ['zombiacı.wav', 'Zombi acı çekme sesi', '108 KB'],
    ['zombivurma.wav', 'Zombi saldırı/vurma sesi', '629 KB'],
    ['insanvurulma.wav', 'Oyuncu vurulma sesi', '1.27 MB'],
])
doc.add_paragraph('Paket İçi Sesler: Silah sesleri (ateş, şarjör, nişan, kılıf) ve karakter yürüme ayak sesleri.')

# ===== 6. 3D MODELLER =====
add_heading('6. 3D Modeller ve Animasyonlar', 1)
add_heading('Zombi Modeli (Assets/zombi/)', 2)
add_table(['Varlık', 'Açıklama', 'Format'], [
    ['Ch10_nonPBR.fbx', 'Ana zombi 3D modeli (Mixamo)', 'FBX (~105 MB)'],
    ['Ch10_nonPBR.prefab', 'Zombi prefab\'ı', 'Prefab'],
    ['ZombieAnim.controller', 'Zombi animasyon kontrolcüsü', 'Controller'],
])

add_heading('Zombi Animasyonları', 2)
add_table(['Animasyon', 'Açıklama'], [
    ['Ch10_nonPBR@Walking.fbx', 'Yürüme animasyonu'],
    ['Ch10_nonPBR@Mutant Swiping.fbx', 'Saldırı (swiping) animasyonu'],
    ['Ch10_nonPBR@Big Hit To Head.fbx', 'Vurulma tepkisi animasyonu'],
    ['Ch10_nonPBR@Zombie Death.fbx', 'Ölüm animasyonu'],
])

add_heading('Silah ve Karakter Modelleri (Paket İçi)', 2)
doc.add_paragraph('AR (Assault Rifle) – Tüfek modeli + animasyonlar', style='List Bullet')
doc.add_paragraph('Handgun – Tabanca modeli + animasyonlar', style='List Bullet')
doc.add_paragraph('First Person karakter kol/el modelleri', style='List Bullet')

# ===== 7. SAHNE =====
add_heading('7. Sahne (Scene) Yapısı', 1)
add_table(['Sahne', 'Konum', 'Boyut'], [
    ['SampleScene.unity', 'Assets/Scenes/', '181 KB'],
])
doc.add_paragraph('Sahne İçeriği:', style='List Bullet')
doc.add_paragraph('3 adet Terrain dosyası ile arazi tasarımı', style='List Bullet')
doc.add_paragraph('NavMesh – Zombi AI yol bulma haritası', style='List Bullet')
doc.add_paragraph('Real Stars Skybox – Gece gökyüzü atmosferi', style='List Bullet')
doc.add_paragraph('NatureStarterKit2 ile ağaçlar, çimen ve yer kaplamaları', style='List Bullet')

# ===== 8. PREFABLAR =====
add_heading('8. Prefab\'lar', 1)
add_table(['Kategori', 'Prefab\'lar'], [
    ['Zombi (Özel)', 'Ch10_nonPBR.prefab'],
    ['Karakter', 'P_LPSP_FP_CH.prefab (FPS karakter)'],
    ['Silahlar', 'P_LPSP_WEP_AR_01 (Tüfek), P_LPSP_WEP_Handgun_03 (Tabanca)'],
    ['Mermiler', 'P_LPSP_PROJ_Bullet_01'],
    ['Kovanlar', 'P_LPSP_Casing_Big, P_LPSP_Casing_Small'],
    ['Hasar Alabilir', 'Varil, Gaz Tankı, Hedef'],
    ['Arayüz', 'P_LPSP_UI_Canvas (HUD)'],
    ['Post Process', 'P_LPSP_PP'],
])

# ===== 9. UNITY PAKETLERI =====
add_heading('9. Unity Paketleri (Packages)', 1)
add_table(['Paket', 'Versiyon', 'Kullanım Amacı'], [
    ['Universal Render Pipeline', '14.0.12', 'URP render pipeline'],
    ['Input System', '1.14.0', 'Yeni input sistemi'],
    ['AI Navigation', '1.1.7', 'NavMesh – Zombi yol bulma'],
    ['Cinemachine', '2.10.3', 'Gelişmiş kamera sistemi'],
    ['TextMeshPro', '3.0.7', 'Gelişmiş UI metin'],
    ['Post Processing', '3.4.0', 'Görsel efektler'],
    ['Shader Graph', '14.0.12', 'Shader tasarımı'],
    ['Burst', '1.8.21', 'Performans optimizasyonu'],
    ['Timeline', '1.7.7', 'Cutscene / zaman çizelgesi'],
])

# ===== 10. TERRAIN =====
add_heading('10. Terrain ve Çevre Tasarımı', 1)
add_table(['Dosya', 'Boyut'], [
    ['New Terrain.asset', '544 KB'],
    ['New Terrain 1.asset', '544 KB'],
    ['New Terrain 2.asset', '1.87 MB'],
])
doc.add_paragraph('Terrain Katmanları: layer_ground01, layer_ground02, layer_ground03')
doc.add_paragraph('NatureStarterKit2: Ağaç, çalı, taş mesh\'leri, materyaller ve tekstürler')

# ===== 11. OYUN MEKANIKLERI =====
add_heading('11. Oyun Mekanikleri Özeti', 1)

doc.add_paragraph('Oyun Akışı:', style='List Bullet')
doc.add_paragraph('1. Oyun başlar → ZombieSpawner aktif olur', style='List Bullet 2')
doc.add_paragraph('2. Zombiler spawn olur (max 5 eşzamanlı, 3sn aralık)', style='List Bullet 2')
doc.add_paragraph('3. Zombiler NavMesh ile oyuncuya doğru hareket eder', style='List Bullet 2')
doc.add_paragraph('4. Menzile girince saldırır (10 hasar)', style='List Bullet 2')
doc.add_paragraph('5. Oyuncu ateş eder → Raycast ile vuruş → Zombi hasar alır', style='List Bullet 2')
doc.add_paragraph('6. 15 zombi öldürülünce → KAZANDIN!', style='List Bullet 2')
doc.add_paragraph('7. Oyuncu canı 0 olursa → Ölüm sekansı + menü', style='List Bullet 2')

add_heading('Kontroller', 2)
add_table(['Tuş', 'İşlev'], [
    ['WASD', 'Hareket'],
    ['Mouse', 'Bakış (kamera)'],
    ['Sol Tık', 'Ateş etme'],
    ['R', 'Şarjör değiştirme'],
    ['F', 'El feneri aç/kapa'],
])

# ===== 12. GENEL DEGERLENDIRME =====
add_heading('12. Genel Değerlendirme', 1)

add_heading('Tamamlanan Özellikler', 2)
completed = [
    'FPS karakter kontrolü (hareket + kamera)',
    'Silah sistemi (ateş etme, şarjör, animasyon, ses)',
    'Zombi AI sistemi (NavMesh, saldırı, vurulma, ölüm)',
    'Zombi spawn sistemi (ayarlanabilir, sürekli)',
    'Oyuncu sağlık sistemi (can barı, hasar, dokunulmazlık)',
    'Kan ekranı efekti (blood screen overlay)',
    'Kalp atışı sistemi (düşük canda)',
    'Ölüm sekansı (kamera düşme + ölüm menüsü)',
    'Kazanma sistemi (15 zombi = bölüm geçme)',
    'UI sistemi (HUD, mermi sayacı, hedef sayacı)',
    'El feneri sistemi',
    'Ses sistemi (silah, zombi, çevre, oyuncu)',
    'Gece atmosferi (yıldızlı skybox)',
    'Doğa ortamı (terrain, ağaçlar, çimen)',
    'NavMesh yapılandırması',
]
for item in completed:
    doc.add_paragraph('✅ ' + item, style='List Bullet')

add_heading('İstatistikler', 2)
add_table(['Metrik', 'Değer'], [
    ['Toplam Özel Script Sayısı', '7 adet'],
    ['Toplam Özel Kod Satırı', '~708 satır'],
    ['Paket Script Sayısı', '~35+ adet'],
    ['Özel Ses Dosyası', '5 adet'],
    ['Zombi Animasyonu', '4 adet'],
    ['Terrain Dosyası', '3 adet'],
    ['Silah Modeli', '2 adet (AR + Tabanca)'],
    ['Asset Paketi', '5 adet'],
])

# Save
output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'Proje_Ilerleme_Raporu.docx')
doc.save(output_path)
print(f"Rapor kaydedildi: {output_path}")
